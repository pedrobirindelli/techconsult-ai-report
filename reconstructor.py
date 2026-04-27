import os
import json
import requests
import tempfile
import docx
from docx.shared import Inches

TEMPLATE_DOCX = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\Relatórios Modelos\Ed MOnt Blanc rel -.docx"
OUTPUT_JSON = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\output_report.json"
FINAL_DOCX = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\Relatorio_Final_IA.docx"
TEMP_MEDIA = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\temp_media"

def download_image_temporarily(url):
    """Downloads an image from a URL to a temporary file."""
    try:
        # Check if the URL matches one of our local files
        filename = os.path.basename(url.split('?')[0])
        local_path = os.path.join(TEMP_MEDIA, filename)
        if os.path.exists(local_path):
            return local_path
            
        # If not, download it
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            fd, path = tempfile.mkstemp(suffix=".jpg")
            with os.fdopen(fd, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            return path
    except Exception as e:
        print(f"Failed to fetch image {url}: {e}")
    return None

def clear_document_body(doc):
    """Removes all paragraphs and tables from the body, keeping headers/footers/styles."""
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)

def apply_style(doc, element_type, text):
    """Tries to apply specific styles based on what exists in the document."""
    # This tries to match common Word styles or defaults to Normal
    style_map = {
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "paragraph": "Normal"
    }
    target_style = style_map.get(element_type, "Normal")
    
    # We create the paragraph
    p = doc.add_paragraph(text)
    
    try:
        p.style = doc.styles[target_style]
    except KeyError:
        # If the exact style doesn't exist in the old doc, fallback to Normal
        pass

def reconstruct_word():
    print("Iniciando a reconstrução do Relatório em Word...")
    
    # 1. Open the original document to inherit styles
    doc = docx.Document(TEMPLATE_DOCX)
    
    # 2. Delete old content
    print("Limpando dados antigos...")
    clear_document_body(doc)
    
    # 3. Read the AI structured output
    with open(OUTPUT_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)
        
    print("Injetando nova estrutura clonada...")
    for block in data:
        b_type = block.get("type", "")
        
        if b_type.startswith("heading") or b_type == "paragraph":
            apply_style(doc, b_type, block.get("text", ""))
            
        elif b_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            
            # Create table
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid' # Standard Word table style
            
            # Fill headers
            hdr_cells = table.rows[0].cells
            for idx, h in enumerate(headers):
                hdr_cells[idx].text = str(h)
                
            # Fill rows
            for i, row in enumerate(rows):
                row_cells = table.rows[i+1].cells
                for j, cell_val in enumerate(row):
                    if j < len(row_cells):
                        row_cells[j].text = str(cell_val)
                        
            # Add a small spacing after table
            doc.add_paragraph("")
            
        elif b_type == "image":
            img_url = block.get("url", "")
            caption = block.get("caption", "")
            
            if img_url:
                local_img = download_image_temporarily(img_url)
                if local_img:
                    try:
                        # Insert image centered with 5 inches width
                        p = doc.add_paragraph()
                        p.alignment = 1 # Center alignment
                        run = p.add_run()
                        run.add_picture(local_img, width=Inches(5.0))
                    except Exception as e:
                        print(f"Could not insert image: {e}")
                        
            # Insert caption
            if caption:
                p_caption = doc.add_paragraph(caption)
                p_caption.alignment = 1 # Center
                try:
                    p_caption.style = doc.styles["Caption"]
                except KeyError:
                    pass

    # 4. Save the final document
    doc.save(FINAL_DOCX)
    print(f"\nSUCESSO! Relatório finalizado e salvo em: {FINAL_DOCX}")
    
    # Try to open the file automatically in Windows
    try:
        os.startfile(FINAL_DOCX)
    except:
        pass

if __name__ == "__main__":
    reconstruct_word()
