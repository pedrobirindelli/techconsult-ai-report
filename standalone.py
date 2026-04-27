import sys
import os
import time
import json
import uuid
import docx
import pandas as pd
import requests
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from google import genai
from dotenv import load_dotenv
import re
from urllib.parse import urlparse
from docx.shared import Inches
from bs4 import BeautifulSoup
import PyPDF2
import openpyxl
import webbrowser
import threading
import logging
from logging.handlers import RotatingFileHandler

# Configuração de Logs
def setup_logging():
    log_file = os.path.join(os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else __file__), 'debug.log')
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            RotatingFileHandler(log_file, maxBytes=1024*1024, backupCount=2),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger('TechConsult')

logger = setup_logging()

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

dist_folder = resource_path('dist')
app = Flask(__name__, static_folder=dist_folder, static_url_path='/')
CORS(app)

# Force load dotenv from the bundled resources
load_dotenv(resource_path(".env"))

api_key = os.getenv("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)

# Folder for uploads and rules in the same folder as the executable
base_run_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else __file__)
UPLOAD_FOLDER = os.path.join(base_run_dir, 'uploads_standalone')
RULES_FILE = os.path.join(base_run_dir, 'rules.json')
PROJECTS_FOLDER = os.path.join(base_run_dir, 'projects')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROJECTS_FOLDER, exist_ok=True)

def find_project_file(project_name, original_project_name, category, filename):
    """Busca um arquivo no projeto atual ou no original (caso de renomeação)."""
    if project_name:
        path = os.path.join(PROJECTS_FOLDER, project_name, category, filename)
        if os.path.exists(path): return path
    if original_project_name and original_project_name != project_name:
        path = os.path.join(PROJECTS_FOLDER, original_project_name, category, filename)
        if os.path.exists(path): return path
    return None

def load_persistent_rules():
    if os.path.exists(RULES_FILE):
        try:
            with open(RULES_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except: return []
    return []

def verify_environment():
    """Verifica se todos os componentes necessários existem e funcionam."""
    issues = []
    logger.info("Iniciando verificação de ambiente...")
    
    # 1. Verificar pasta dist (Interface)
    if not os.path.exists(dist_folder):
        msg = f"Erro Crítico: Pasta de interface '{dist_folder}' não encontrada."
        logger.error(msg)
        issues.append(msg)
    
    # 2. Verificar .env e API Key
    env_path = resource_path(".env")
    if not os.path.exists(env_path):
        logger.warning("Arquivo .env não encontrado.")
        issues.append("Aviso: Arquivo .env não encontrado na raiz.")
    else:
        load_dotenv(env_path)
        if not os.getenv("GEMINI_API_KEY"):
            logger.error("GEMINI_API_KEY ausente no .env")
            issues.append("Erro: Chave GEMINI_API_KEY não encontrada no arquivo .env.")
        else:
            logger.info("Chave API detectada.")

    # 3. Testar dependências críticas
    try:
        import pandas as pd
        import openpyxl
        logger.info(f"Pandas e Openpyxl carregados com sucesso (Pandas v{pd.__version__}).")
    except Exception as e:
        msg = f"Erro de Dependência (Excel): Não foi possível carregar as bibliotecas de dados ({e})."
        logger.error(msg)
        issues.append(msg)

    try:
        import docx
        doc_test = docx.Document()
        logger.info("Python-Docx funcionando.")
    except Exception as e:
        msg = f"Erro de Dependência (Word): Falha ao carregar motor de relatórios ({e})."
        logger.error(msg)
        issues.append(msg)

    # 4. Verificar conectividade básica (Opcional, não bloqueante)
    try:
        requests.get("https://generativelanguage.googleapis.com", timeout=5)
        logger.info("Conexão com Google Gemini API disponível.")
    except Exception as e:
        logger.warning(f"Possível problema de rede/firewall: {e}")
        issues.append("Aviso: Não foi possível validar conexão com API do Google. Verifique seu Firewall.")

    # 5. Verificar permissões de escrita
    try:
        test_file = os.path.join(UPLOAD_FOLDER, 'test_write.txt')
        with open(test_file, 'w') as f:
            f.write('test')
        os.remove(test_file)
        logger.info("Permissões de escrita OK.")
    except Exception as e:
        msg = f"Erro de Permissão: Sem acesso de escrita na pasta uploads ({e})."
        logger.error(msg)
        issues.append(msg)
        
    return issues

def save_persistent_rules(rules):
    with open(RULES_FILE, 'w', encoding='utf-8') as f:
        json.dump(rules, f, ensure_ascii=False, indent=2)

# Helper functions
def extract_urls(text):
    if pd.isna(text): return []
    return re.findall(r'(https?://[^\s\]\n]+)', str(text))

def download_file(url, folder):
    try:
        parsed = urlparse(url)
        filename = os.path.basename(parsed.path) or f"{uuid.uuid4().hex}.bin"
        filepath = os.path.join(folder, filename)
        if os.path.exists(filepath): return filepath
        response = requests.get(url, stream=True, timeout=15)
        if response.status_code == 200:
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(8192):
                    f.write(chunk)
            return filepath
    except Exception as e:
        print(f"Failed download {url}: {e}")
    return None

def extract_text_from_docx(file_path):
    try:
        if not file_path.lower().endswith('.docx'): return ""
        doc = docx.Document(file_path)
        fullText = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                fullText.append(' | '.join(row_data))
        return '\n'.join(fullText)
    except: return ""

def extract_text_from_pdf(file_path):
    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except: return ""

def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith('.docx'): return extract_text_from_docx(file_path)
    if ext.endswith('.pdf'): return extract_text_from_pdf(file_path)
    return ""

def clear_document_body(doc):
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)

def reconstruct_doc(template_path, json_data, url_to_local, output_path):
    try:
        doc = docx.Document(template_path)
    except:
        doc = docx.Document()
        
    clear_document_body(doc)
    style_map = {"heading1": "Heading 1", "heading2": "Heading 2", "heading3": "Heading 3", "paragraph": "Normal"}
    
    for block in json_data:
        if not isinstance(block, dict):
            logger.warning(f"Ignorando bloco de dado inválido (não é um objeto): {block}")
            continue
            
        b_type = block.get("type", "")
        if b_type.startswith("heading") or b_type == "paragraph":
            p = doc.add_paragraph(block.get("text", "") or block.get("content", ""))
            try: p.style = doc.styles[style_map.get(b_type, "Normal")]
            except: pass
        elif b_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Table Grid'
                for idx, h in enumerate(headers): table.rows[0].cells[idx].text = str(h)
                for i, row in enumerate(rows):
                    for j, val in enumerate(row): table.rows[i+1].cells[j].text = str(val)
        elif b_type == "image":
            img_url = block.get("url", "")
            local_img = url_to_local.get(img_url)
            if local_img:
                try:
                    p = doc.add_paragraph()
                    p.alignment = 1
                    p.add_run().add_picture(local_img, width=Inches(5.0))
                except: pass
            caption = block.get("caption", "")
            if caption:
                p_cap = doc.add_paragraph(caption)
                p_cap.alignment = 1
                try: p_cap.style = doc.styles["Caption"]
                except: pass
    doc.save(output_path)

@app.route('/')
def index():
    return send_from_directory(dist_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    if os.path.exists(os.path.join(dist_folder, path)):
        return send_from_directory(dist_folder, path)
    return send_from_directory(dist_folder, 'index.html')

@app.route('/api/rules', methods=['GET'])
def list_rules():
    return jsonify(load_persistent_rules())

@app.route('/api/rules', methods=['POST'])
def add_rule():
    data = request.json
    rules = load_persistent_rules()
    new_rule = {"id": uuid.uuid4().hex, "text": data.get("text", ""), "active": True}
    rules.append(new_rule)
    save_persistent_rules(rules)
    return jsonify(new_rule)

@app.route('/api/rules/<rule_id>', methods=['DELETE'])
def delete_rule(rule_id):
    rules = load_persistent_rules()
    rules = [r for r in rules if r['id'] != rule_id]
    save_persistent_rules(rules)
    return jsonify({"success": True})
@app.route('/api/rules/<rule_id>', methods=['PUT'])
def update_rule(rule_id):
    data = request.json
    rules = load_persistent_rules()
    for r in rules:
        if r['id'] == rule_id:
            if 'text' in data: r['text'] = data['text']
            if 'active' in data: r['active'] = data['active']
            break
    save_persistent_rules(rules)
    return jsonify({"success": True})

@app.route('/api/projects', methods=['GET'])
def list_projects():
    if not os.path.exists(PROJECTS_FOLDER): return jsonify([])
    projects = []
    for d in os.listdir(PROJECTS_FOLDER):
        p_path = os.path.join(PROJECTS_FOLDER, d)
        if os.path.isdir(p_path):
            projects.append({
                "name": d,
                "created_at": time.ctime(os.path.getctime(p_path))
            })
    return jsonify(projects)

@app.route('/api/projects/save', methods=['POST'])
def save_project():
    try:
        import shutil
        data = request.form
        project_name = data.get('project_name', 'Sem Nome').strip()
        if not project_name: return jsonify({"error": "Nome inválido"}), 400
        
        project_dir = os.path.join(PROJECTS_FOLDER, project_name)
        os.makedirs(project_dir, exist_ok=True)
        
        original_project_name = data.get('original_project_name')
        
        mapping = {
            'excel_files': ('excel', 'remote_excel'),
            'template_files': ('template', 'remote_template'),
            'source_files': ('source', 'remote_source'),
            'visual_files': ('visual', 'remote_visual')
        }
        
        for key, (folder, remote_key) in mapping.items():
            folder_path = os.path.join(project_dir, folder)
            os.makedirs(folder_path, exist_ok=True)
            
            files = request.files.getlist(key)
            for f in files:
                if f.filename:
                    f.save(os.path.join(folder_path, f.filename))
            
            if original_project_name and original_project_name != project_name:
                remote_data = data.get(remote_key, '[]')
                try:
                    remote_names = json.loads(remote_data)
                    if isinstance(remote_names, str): remote_names = [remote_names]
                except:
                    remote_names = [remote_data] if remote_data and remote_data != '[]' else []
                
                src_folder = os.path.join(PROJECTS_FOLDER, original_project_name, folder)
                for f_name in remote_names:
                    src_path = os.path.join(src_folder, f_name)
                    dst_path = os.path.join(folder_path, f_name)
                    if os.path.exists(src_path) and not os.path.exists(dst_path):
                        shutil.copy2(src_path, dst_path)
        
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/projects/load/<name>', methods=['GET'])
def load_project(name):
    project_dir = os.path.join(PROJECTS_FOLDER, name)
    if not os.path.exists(project_dir): return jsonify({"error": "Projeto não encontrado"}), 404
    
    result = {"excel": [], "template": [], "source": [], "visual": [], "reports": []}
    for cat in result.keys():
        cat_dir = os.path.join(project_dir, cat)
        if os.path.exists(cat_dir):
            result[cat] = sorted(os.listdir(cat_dir), reverse=True)
            
    return jsonify(result)

@app.route('/api/projects/<name>', methods=['DELETE'])
def delete_project(name):
    try:
        import shutil, stat
        from urllib.parse import unquote
        project_dir = os.path.join(PROJECTS_FOLDER, unquote(name))
        if os.path.exists(project_dir):
            def on_rm_error(func, path, _):
                os.chmod(path, stat.S_IWRITE)
                func(path)
            shutil.rmtree(project_dir, onerror=on_rm_error)
            return jsonify({"success": True})
        return jsonify({"error": "Projeto não encontrado"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/projects/download_report/<project_name>/<report_name>')
def download_project_report(project_name, report_name):
    report_dir = os.path.join(PROJECTS_FOLDER, project_name, 'reports')
    return send_from_directory(report_dir, report_name, as_attachment=True)

@app.route('/api/status', methods=['GET'])
def get_status():
    issues = verify_environment()
    return jsonify({
        "status": "ok" if not issues else "warning",
        "issues": issues,
        "api_key_configured": bool(os.getenv("GEMINI_API_KEY")),
        "version": "1.3-portable"
    })

@app.route('/estimate', methods=['POST'])
def estimate_cost():
    try:
        print("Iniciando estimativa de tokens...")
        excel_files = request.files.getlist('excel_files')
        template_files = request.files.getlist('template_files')
        source_files = request.files.getlist('source_files')
        
        project_name = request.form.get('project_name')
        original_project_name = request.form.get('original_project_name')
        remote_excel = json.loads(request.form.get('remote_excel', '[]'))
        remote_template = json.loads(request.form.get('remote_template', '[]'))
        remote_source = json.loads(request.form.get('remote_source', '[]'))

        media_count = 0
        text_size_bytes = 0
        run_id = uuid.uuid4().hex
        run_folder = os.path.join(UPLOAD_FOLDER, "est_" + run_id)
        os.makedirs(run_folder, exist_ok=True)
        
        # Process Remote Excels
        if remote_excel:
            for f_name in remote_excel:
                f_path = find_project_file(project_name, original_project_name, 'excel', f_name)
                if f_path:
                    try:
                        df = pd.read_excel(f_path)
                        for col in df.columns:
                            col_data = df[col].dropna().astype(str)
                            for val in col_data:
                                if 'http' in val:
                                    urls = extract_urls(val)
                                    media_count += len([u for u in urls if 'supabase.co' in u])
                    except: pass

        for idx, excel in enumerate(excel_files):
            if excel.filename:
                print(f"Processando planilha: {excel.filename}")
                path = os.path.join(run_folder, f"est_data_{idx}.xlsx")
                excel.save(path)
                try:
                    df = pd.read_excel(path)
                    # Otimização: processar apenas colunas que costumam ter URLs
                    for col in df.columns:
                        col_data = df[col].dropna().astype(str)
                        for val in col_data:
                            if 'http' in val:
                                urls = extract_urls(val)
                                media_count += len([u for u in urls if 'supabase.co' in u])
                except Exception as ex:
                    print(f"Erro ao ler excel {excel.filename}: {ex}")

        # Process Remote Templates
        if remote_template:
            for f_name in remote_template:
                f_path = find_project_file(project_name, original_project_name, 'template', f_name)
                if f_path:
                    try:
                        text_content = extract_text(f_path)
                        text_size_bytes += len(text_content)
                    except: pass

        for t in template_files:
            if t.filename:
                print(f"Lendo template: {t.filename}")
                path = os.path.join(run_folder, t.filename)
                t.save(path)
                try:
                    text_content = extract_text(path)
                    text_size_bytes += len(text_content)
                except Exception as ex:
                    print(f"Erro ao extrair texto de {t.filename}: {ex}")

        # Process Remote Sources
        if remote_source:
            for f_name in remote_source:
                f_path = find_project_file(project_name, original_project_name, 'source', f_name)
                if f_path:
                    try:
                        text_content = extract_text(f_path)
                        text_size_bytes += len(text_content)
                    except: pass
                
        for s in source_files:
            if s.filename:
                print(f"Lendo fonte: {s.filename}")
                path = os.path.join(run_folder, s.filename)
                s.save(path)
                try:
                    text_content = extract_text(path)
                    text_size_bytes += len(text_content)
                except Exception as ex:
                    print(f"Erro ao extrair texto de {s.filename}: {ex}")

        text_tokens = int(text_size_bytes / 4)
        total_tokens = text_tokens + (media_count * 258)
        print(f"Estimativa concluída: {total_tokens} tokens.")
        return jsonify({
            "media_count": media_count, 
            "text_tokens": text_tokens, 
            "total_tokens": total_tokens, 
            "estimated_usd": (total_tokens/1000000)*1.25
        })
    except PermissionError: 
        logger.error("Erro de permissão: arquivos abertos.")
        return jsonify({"error": "Arquivos Excel ou Word estão abertos em outro programa. Feche-os e tente novamente."}), 500
    except Exception as e: 
        logger.exception("Erro fatal na estimativa de tokens")
        return jsonify({"error": f"Erro interno ao calcular tokens. Verifique o arquivo debug.log para detalhes. Erro: {str(e)}"}), 500

@app.route('/generate', methods=['POST'])
def generate_report():
    try:
        excel_files = request.files.getlist('excel_files')
        template_files = request.files.getlist('template_files')
        source_files = request.files.getlist('source_files')
        rules = request.form.get('knowledge_rules', '')
        
        project_name = request.form.get('project_name')
        original_project_name = request.form.get('original_project_name')
        remote_excel = json.loads(request.form.get('remote_excel', '[]'))
        remote_template = json.loads(request.form.get('remote_template', '[]'))
        remote_source = json.loads(request.form.get('remote_source', '[]'))
        remote_visual = request.form.get('remote_visual')
        
        # Validação baseada em local ou remoto
        has_excel = len(excel_files) > 0 or len(remote_excel) > 0
        has_template = len(template_files) > 0 or len(remote_template) > 0
        
        if not has_excel or not has_template:
            return jsonify({"error": "Faltam arquivos obrigatórios (Planilha ou Template)"}), 400
            
        run_folder = os.path.join(UPLOAD_FOLDER, uuid.uuid4().hex)
        os.makedirs(run_folder, exist_ok=True)
        
        all_excel_str = ""
        url_to_local = {}
        
        # Process Remote Excels
        if remote_excel:
            for f_name in remote_excel:
                f_path = find_project_file(project_name, original_project_name, 'excel', f_name)
                if f_path:
                    try:
                        df = pd.read_excel(f_path)
                        all_excel_str += f"\n[Arquivo Projeto: {f_name}]\n" + df.to_json(orient="records", force_ascii=False)
                        for index, row in df.iterrows():
                            for col in df.columns:
                                val = row[col]
                                if isinstance(val, str):
                                    for url in extract_urls(val):
                                        if 'supabase.co' in url:
                                            local = download_file(url, run_folder)
                                            if local: url_to_local[url] = local
                    except: pass

        for idx, excel in enumerate(excel_files):
            if excel.filename:
                path = os.path.join(run_folder, f"data_{idx}.xlsx")
                excel.save(path)
                df = pd.read_excel(path)
                all_excel_str += f"\n[Planilha {idx+1}]\n" + df.to_json(orient="records", force_ascii=False)
                for index, row in df.iterrows():
                    for col in df.columns:
                        for url in extract_urls(row[col]):
                            if 'supabase.co' in url:
                                local = download_file(url, run_folder)
                                if local: url_to_local[url] = local

        gemini_files = []
        media_mapping_text = "\nMAPEAMENTO DE IMAGENS:\n"
        local_to_gemini = {}
        for url, local_path in url_to_local.items():
            if local_path.endswith(('.jpeg', '.jpg', '.png', '.webp')):
                if local_path not in local_to_gemini:
                    g_file = client.files.upload(file=local_path)
                    while g_file.state.name == "PROCESSING": time.sleep(1); g_file = client.files.get(name=g_file.name)
                    local_to_gemini[local_path] = g_file
                    gemini_files.append(g_file)
                media_mapping_text += f"- {os.path.basename(local_path)} -> {url}\n"

        all_old_texts = ""; base_template_path = None
        
        # Process Remote Templates
        if remote_template:
            for f_name in remote_template:
                f_path = find_project_file(project_name, original_project_name, 'template', f_name)
                if f_path:
                    if not base_template_path and f_name.lower().endswith('.docx'):
                        base_template_path = f_path
                    all_old_texts += f"\n[Referência Projeto: {f_name}]\n" + extract_text(f_path)

        for idx, t in enumerate(template_files):
            if t.filename:
                path = os.path.join(run_folder, f"template_{idx}_{t.filename}")
                t.save(path)
                if not base_template_path and path.lower().endswith('.docx'): base_template_path = path
                all_old_texts += f"\n[Referência {idx+1}]\n" + extract_text(path)

        sources_text = ""
        # Process Remote Sources
        if remote_source:
            for f_name in remote_source:
                f_path = find_project_file(project_name, original_project_name, 'source', f_name)
                if f_path:
                    sources_text += f"\n[Fonte Projeto: {f_name}]\n" + extract_text(f_path)

        for idx, s in enumerate(source_files):
            if s.filename:
                path = os.path.join(run_folder, f"source_{idx}_{s.filename}")
                s.save(path)
                sources_text += f"\n[Fonte {idx+1}]\n" + extract_text(path)

        sys_inst = f"Você é um Perito... {media_mapping_text}"
        user_prompt = f"Gere o laudo... {all_excel_str}\n{all_old_texts}\n{sources_text}\nREGRAS: {rules}"
        
        response = client.models.generate_content(
            model='gemini-2.5-pro', contents=gemini_files + [user_prompt],
            config=genai.types.GenerateContentConfig(system_instruction=sys_inst, temperature=0.2, response_mime_type="application/json")
        )
        
        raw = response.text.strip()
        if raw.startswith("```json"): raw = raw[7:-3].strip()
        elif raw.startswith("```"): raw = raw[3:-3].strip()
        
        final_doc = os.path.join(run_folder, "Resultado.docx")
        reconstruct_doc(base_template_path, json.loads(raw), url_to_local, final_doc)
        return send_file(final_doc, as_attachment=True, download_name="Relatorio_IA.docx")
        
    except PermissionError: return jsonify({"error": "arquivos abertos por outros aplicativos"}), 500
    except Exception as e: 
        logger.exception("Erro fatal na geração")
        return jsonify({"error": str(e)}), 500

def open_browser():
    time.sleep(2)
    webbrowser.open("http://127.0.0.1:4567")

if __name__ == '__main__':
    print("="*50)
    print("TECHCONSULT AI - INICIANDO SISTEMA")
    print("="*50)
    
    issues = verify_environment()
    if issues:
        print("\nDIAGNÓSTICOS DE INICIALIZAÇÃO:")
        for issue in issues:
            print(f" [!] {issue}")
        print("\nO programa tentará rodar mesmo assim, mas pode apresentar falhas.")
    else:
        print("\n[OK] Ambiente verificado e pronto para uso.")
    
    print(f"\nInterface: http://127.0.0.1:4567")
    print("Mantenha esta janela aberta enquanto utiliza o sistema.")
    print("="*50 + "\n")

    if os.getenv("GEMINI_API_KEY"): 
        threading.Thread(target=open_browser, daemon=True).start()
    else:
        print("\n[AVISO] Chave API não configurada. O navegador não abrirá automaticamente.")
        print("Configure o arquivo .env com a chave GEMINI_API_KEY.\n")
        
    app.run(port=4567, debug=False)
