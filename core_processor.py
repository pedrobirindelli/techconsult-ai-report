import os
import time
import json
import docx
from google import genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=api_key)

DOWNLOAD_DIR = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\temp_media"
EXCEL_FILE = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\Export do excel_dados da vistoria\Dados_Wayne Barra ap206 t1_25-04-2026.xlsx"
TEMPLATE_DOCX = r"C:\Users\pedro\OneDrive\01 - Pessoal\07 - Projetos Pessoais\12 - Programas\techconsult-ai-report\Relatórios Modelos\Ed MOnt Blanc rel -.docx"

def extract_text_from_docx(file_path):
    print("Extracting text from reference Word document...")
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        if para.text.strip():
            fullText.append(para.text)
    
    # Also grab tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            fullText.append(' | '.join(row_data))
            
    return '\n'.join(fullText)

def upload_to_gemini(file_path, mime_type=None):
    """Uploads the given file to Gemini."""
    print(f"Uploading '{os.path.basename(file_path)}' to Gemini...")
    file = client.files.upload(file=file_path)
    
    # Wait for processing if it's an audio or video
    while file.state.name == "PROCESSING":
        print("Waiting for file processing...", end="\r")
        time.sleep(2)
        file = client.files.get(name=file.name)
        
    if file.state.name == "FAILED":
        raise ValueError(f"File processing failed for {file_path}")
        
    print(f"Uploaded successfully as {file.name}")
    return file

def generate_report():
    print("--- Início do Processamento de Inteligência Artificial ---")
    
    # 1. Read the old report text
    old_report_text = extract_text_from_docx(TEMPLATE_DOCX)
    
    # 2. Upload media
    uploaded_files = []
    if os.path.exists(DOWNLOAD_DIR):
        for filename in os.listdir(DOWNLOAD_DIR):
            if filename.endswith(".jpeg") or filename.endswith(".jpg") or filename.endswith(".webm") or filename.endswith(".mp4"):
                filepath = os.path.join(DOWNLOAD_DIR, filename)
                try:
                    gemini_file = upload_to_gemini(filepath)
                    uploaded_files.append(gemini_file)
                except Exception as e:
                    print(f"Skipping file {filename} due to error: {e}")
                
    # 3. Read Excel Data
    import pandas as pd
    df = pd.read_excel(EXCEL_FILE)
    excel_data_str = df.to_json(orient="records", force_ascii=False)
    
    # 4. Construct Prompt
    system_instruction = """
    Você é um Engenheiro Civil e Perito Técnico sênior.
    Seu objetivo é gerar um novo relatório de vistoria técnica.
    
    INSTRUÇÕES:
    1. Leia o texto do 'Laudo Antigo de Referência' para aprender o tom de voz, a formatação das tabelas e a estrutura de capítulos.
    2. Leia os dados extraídos do Excel da NOVA vistoria.
    3. Analise as fotos e os áudios fornecidos em anexo. Transcreva mentalmente os áudios e cruze com o que vê nas fotos correspondentes.
    4. Crie o conteúdo do NOVO laudo, mantendo estritamente a mesma estrutura de tópicos do 'Laudo Antigo'. Substitua os dados antigos pelos dados novos encontrados no Excel, nas fotos e nos áudios.
    5. Retorne a sua resposta EXCLUSIVAMENTE em um formato JSON estruturado, onde você define blocos de conteúdo.

    O formato JSON DEVE seguir este padrão:
    [
        {"type": "heading1", "text": "1. Introdução"},
        {"type": "paragraph", "text": "Na data de 25/04/2026, foi realizada a vistoria..."},
        {"type": "image", "caption": "Foto 1: Fissura na parede leste identificada na imagem 0.322317..."},
        {"type": "table", "headers": ["Ambiente", "Condição"], "rows": [["Sala", "Pintura descascando"]]}
    ]
    """
    
    user_prompt = f"""
    --- LAUDO ANTIGO DE REFERÊNCIA (APRENDA A ESTRUTURA COM ISTO) ---
    {old_report_text[:10000]} # Truncating to avoid massive token load for this test, you can adjust later
    
    --- DADOS DA NOVA VISTORIA (PLANILHA EXCEL) ---
    {excel_data_str}
    
    Gere o JSON estruturado do novo relatório. Retorne APENAS o JSON válido, sem formatação de markdown (```json).
    """
    
    print("\nIniciando raciocínio do modelo (isso pode demorar devido à análise das fotos/áudios)...")
    
    # Combine everything for the model
    contents = uploaded_files + [user_prompt]
    
    response = client.models.generate_content(
        model='gemini-2.5-pro',
        contents=contents,
        config=genai.types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.2, # Low temperature for more deterministic/professional output
            response_mime_type="application/json"
        )
    )
    
    print("\n--- Relatório Gerado com Sucesso ---")
    
    # Save the output
    output_path = "output_report.json"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(response.text)
        
    print(f"O esqueleto do relatório foi salvo em: {output_path}")

if __name__ == "__main__":
    generate_report()
