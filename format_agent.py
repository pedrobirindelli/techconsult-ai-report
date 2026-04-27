import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, font_name, size_pt, color_hex, bold, italic):
    if font_name:
        run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        color_hex = color_hex.lstrip('#')
        try:
            run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
        except: pass
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def format_class(doc, text_class, font_name, size_pt, color_hex, bold, italic):
    """
    text_class: 'titulo', 'subtitulo', 'texto', 'legenda'
    """
    for para in doc.paragraphs:
        is_target = False
        if text_class == 'titulo' and para.style.name.startswith('Heading 1'): is_target = True
        elif text_class == 'subtitulo' and para.style.name.startswith('Heading 2'): is_target = True
        elif text_class == 'texto' and para.style.name == 'Normal': is_target = True
        elif text_class == 'legenda' and (para.style.name == 'Caption' or 'figura' in para.text.lower() or 'foto' in para.text.lower()): is_target = True
        
        if is_target:
            for run in para.runs:
                set_font(run, font_name, size_pt, color_hex, bold, italic)

def format_specific_text(doc, search_text, color_hex, bg_color_hex, bold):
    """
    Encontra `search_text` e formata a cor do texto e preenchimento (highlight).
    """
    for para in doc.paragraphs:
        if search_text.lower() in para.text.lower():
            for run in para.runs:
                if search_text.lower() in run.text.lower():
                    set_font(run, None, None, color_hex, bold, None)
                    # bg_color_hex via highlight is limited in python-docx, usually we use run.font.highlight_color
                    # But we can try to set XML shading
                    if bg_color_hex:
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), bg_color_hex.lstrip('#'))
                        run._r.get_or_add_rPr().append(shd)

def align_component(doc, component, alignment):
    """
    component: 'texto', 'imagem', 'tabela', 'titulo'
    alignment: 'left', 'center', 'right', 'justify'
    """
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    wd_align = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    for para in doc.paragraphs:
        if component == 'titulo' and para.style.name.startswith('Heading'):
            para.alignment = wd_align
        elif component == 'texto' and para.style.name == 'Normal':
            para.alignment = wd_align
        elif component == 'imagem':
            # Verifica se tem imagem no parágrafo
            if 'w:drawing' in para._p.xml or 'v:imagedata' in para._p.xml:
                para.alignment = wd_align

    if component == 'tabela':
        for table in doc.tables:
            table.alignment = wd_align

def add_cover(doc, title, subtitle, author):
    doc.paragraphs[0].insert_paragraph_before("\n\n\n\n\n")
    p_title = doc.paragraphs[0].insert_paragraph_before(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.size = Pt(28)
    p_title.runs[0].bold = True

    p_sub = p_title.insert_paragraph_before(subtitle)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.size = Pt(18)

    p_auth = p_sub.insert_paragraph_before(author)
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.runs[0].font.size = Pt(14)
    
    doc.paragraphs[0].insert_paragraph_before("\n\n\n\n\n")
    
    # Page break
    doc.paragraphs[0].insert_paragraph_before("").add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

def insert_toc(doc):
    p = doc.paragraphs[0].insert_paragraph_before("ÍNDICE (Atualize este campo no Word)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    # Inserir o campo XML de TOC
    paragraph = doc.paragraphs[0].insert_paragraph_before("")
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    paragraph.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

def add_signature_block(doc, names):
    """
    names: list of strings (e.g. ['Pedro', 'João'])
    """
    doc.add_paragraph("\n\n\n\n")
    for name in names:
        p = doc.add_paragraph("________________________________________")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph(name)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

def execute_formatting_plan(file_path, plan_steps, output_path):
    """
    plan_steps = [
        {"action": "add_cover", "params": {"title": "Laudo", "subtitle": "Vistoria", "author": "TechConsult"}},
        {"action": "format_class", "params": {"text_class": "titulo", "color_hex": "#0000FF", ...}},
        ...
    ]
    """
    doc = docx.Document(file_path)
    
    for step in plan_steps:
        action = step.get('action')
        params = step.get('params', {})
        
        try:
            if action == 'add_cover':
                add_cover(doc, params.get('title', ''), params.get('subtitle', ''), params.get('author', ''))
            elif action == 'insert_toc':
                insert_toc(doc)
            elif action == 'format_class':
                format_class(doc, params.get('text_class'), params.get('font_name'), params.get('size_pt'), params.get('color_hex'), params.get('bold'), params.get('italic'))
            elif action == 'format_specific_text':
                format_specific_text(doc, params.get('search_text'), params.get('color_hex'), params.get('bg_color_hex'), params.get('bold'))
            elif action == 'align_component':
                align_component(doc, params.get('component'), params.get('alignment'))
            elif action == 'add_signature_block':
                add_signature_block(doc, params.get('names', []))
        except Exception as e:
            print(f"Error applying action {action}: {e}")
            
    doc.save(output_path)
