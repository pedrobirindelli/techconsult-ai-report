import os
import time
import shutil
import json
import uuid
import threading
import queue
import tempfile
import docx
import pandas as pd
import requests
import jwt
import logging
from flask import Flask, request, jsonify, send_file, Response, stream_with_context, make_response
from flask_cors import CORS
from google import genai
from dotenv import load_dotenv
from urllib.parse import urlparse
from docx import Document
from docx.shared import Cm
from functools import wraps
from PIL import Image, ExifTags

load_dotenv()

app = Flask(__name__, static_folder='dist', static_url_path='/')
CORS(app)

@app.route('/')
def serve_index():
    return send_file('dist/index.html')

@app.route('/<path:path>')
def serve_static(path):
    if os.path.exists(os.path.join('dist', path)):
        return send_from_directory('dist', path)
    return send_file('dist/index.html')

from flask import send_from_directory

# Configurações Gemini
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    client = genai.Client(api_key=api_key)

# Configurações Supabase (Auth)
SUPABASE_JWT_SECRET = os.getenv("SUPABASE_JWT_SECRET") # Necessário para produção
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://agfkghlczqpyikphmsog.supabase.co")

UPLOAD_FOLDER = 'temp_uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token:
            return jsonify({"error": "Token de autenticação ausente"}), 401
        
        # Em ambiente de desenvolvimento ou se o segredo não estiver configurado, 
        # apenas verificamos se o token existe. Em produção, você deve configurar o SUPABASE_JWT_SECRET.
        if not SUPABASE_JWT_SECRET:
            # print("Aviso: SUPABASE_JWT_SECRET não configurado. Verificação de token simplificada.")
            if token.startswith("Bearer "): 
                try:
                    t = token.replace("Bearer ", "")
                    request.user = jwt.decode(t, options={"verify_signature": False})
                except:
                    pass
                return f(*args, **kwargs)
            return jsonify({"error": "Token inválido"}), 401

        try:
            token = token.replace("Bearer ", "")
            decoded = jwt.decode(token, SUPABASE_JWT_SECRET, algorithms=["HS256"], audience="authenticated")
            # Injeta o usuário na requisição para uso posterior
            request.user = decoded
            return f(*args, **kwargs)
        except Exception as e:
            return jsonify({"error": f"Token inválido: {str(e)}"}), 401
    return decorated

def require_admin(f):
    @wraps(f)
    @require_auth
    def decorated(*args, **kwargs):
        user = getattr(request, 'user', None)
        if not user or user.get('email', '').lower() != 'pedrobirindelli@gmail.com':
            return jsonify({"error": "Acesso negado. Apenas o administrador pode realizar esta ação."}), 403
        return f(*args, **kwargs)
    return decorated



# Helper functions
def extract_urls(text):
    if pd.isna(text): return []
    import re
    return re.findall(r'(https?://[^\s\]\n]+)', str(text))

def fix_image_orientation(filepath):
    try:
        if not filepath.lower().endswith(('.jpg', '.jpeg', '.png', '.webp')):
            return
        img = Image.open(filepath)
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation] == 'Orientation':
                break
        
        exif = img._getexif()
        if exif is not None:
            orientation_val = exif.get(orientation, None)
            if orientation_val == 3:
                img = img.rotate(180, expand=True)
            elif orientation_val == 6:
                img = img.rotate(270, expand=True)
            elif orientation_val == 8:
                img = img.rotate(90, expand=True)
            
            # Remove EXIF and save
            data = list(img.getdata())
            img_without_exif = Image.new(img.mode, img.size)
            img_without_exif.putdata(data)
            img_without_exif.save(filepath, format=img.format if img.format else 'JPEG')
            img.close()
    except Exception as e:
        print(f"Failed to fix EXIF for {filepath}: {e}")

def download_file(url, folder):
    try:
        parsed = urlparse(url)
        filename = os.path.basename(parsed.path) or f"{uuid.uuid4().hex}.bin"
        filepath = os.path.join(folder, filename)
        if os.path.exists(filepath): return filepath
        response = requests.get(url, stream=True, timeout=15)
        if response.status_code == 200:
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(8192):
                    f.write(chunk)
            fix_image_orientation(filepath)
            return filepath
    except Exception as e:
        print(f"Failed download {url}: {e}")
    return None

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        fullText = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                fullText.append(' | '.join(row_data))
        return '\n'.join(fullText)
    except: return ""

def extract_text_from_pdf(file_path):
    try:
        import PyPDF2
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except: return ""

def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith('.docx'): return extract_text_from_docx(file_path)
    elif ext.endswith('.pdf'): return extract_text_from_pdf(file_path)
    return ""

def clear_document_body(doc):
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)

def get_best_style(doc, style_name, default="Normal"):
    try:
        if style_name in doc.styles: return doc.styles[style_name]
    except: pass
    return doc.styles[default]

def copy_paragraph_format(source_p, target_p):
    target_fmt = target_p.paragraph_format
    source_fmt = source_p.paragraph_format
    try:
        target_fmt.alignment = source_fmt.alignment
        target_fmt.first_line_indent = source_fmt.first_line_indent
        target_fmt.left_indent = source_fmt.left_indent
        target_fmt.right_indent = source_fmt.right_indent
        target_fmt.space_before = source_fmt.space_before
        target_fmt.space_after = source_fmt.space_after
        target_fmt.line_spacing = source_fmt.line_spacing
    except: pass

def copy_run_format(source_run, target_run):
    try:
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        if source_run.font.color.rgb: target_run.font.color.rgb = source_run.font.color.rgb
    except: pass

def get_style_dna(doc):
    dna = {}
    for p in doc.paragraphs:
        style_name = p.style.name
        if style_name not in dna:
            model_run = p.runs[0] if p.runs else None
            dna[style_name] = {"p": p, "run": model_run}
    return dna

def add_picture_smart(run, local_img, in_table=False):
    try:
        img = Image.open(local_img)
        w, h = img.size
        img.close()
        if in_table:
            run.add_picture(local_img, width=Cm(7.5))
        else:
            if w > h:
                run.add_picture(local_img, width=Cm(15))
            else:
                run.add_picture(local_img, height=Cm(10))
    except:
        run.add_picture(local_img, width=Cm(10))

def reconstruct_doc(json_data, template_path, output_path, temp_folder):
    try:
        doc = docx.Document(template_path) if template_path and os.path.exists(template_path) else docx.Document()
    except: doc = docx.Document()
        
    style_dna = get_style_dna(doc)
    clear_document_body(doc)
    
    style_map = {"heading1": "Heading 1", "heading2": "Heading 2", "heading3": "Heading 3", "paragraph": "Normal"}
    
    blocks = []
    if isinstance(json_data, list):
        blocks = json_data
    elif isinstance(json_data, dict):
        for key, val in json_data.items():
            if isinstance(val, list):
                blocks = val
                break
        if not blocks:
            blocks = [json_data]
            
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if not isinstance(block, dict): 
            i += 1
            continue
            
        b_type = block.get("type", "")
        
        if b_type.startswith("heading") or b_type == "paragraph":
            text = block.get("text", "")
            if not text.strip(): 
                i += 1
                continue
            p = doc.add_paragraph()
            style_name = style_map.get(b_type, "Normal")
            actual_style_name = get_best_style(doc, style_name).name
            p.style = actual_style_name
            if actual_style_name in style_dna:
                copy_paragraph_format(style_dna[actual_style_name]["p"], p)
            run = p.add_run(text)
            if actual_style_name in style_dna and style_dna[actual_style_name]["run"]:
                copy_run_format(style_dna[actual_style_name]["run"], run)

        elif b_type == "table":
            headers, rows = block.get("headers", []), block.get("rows", [])
            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                try: table.style = 'Table Grid'
                except: pass
                for idx, h in enumerate(headers):
                    cell = table.rows[0].cells[idx]
                    cell.text = str(h)
                    for para in cell.paragraphs:
                        para.alignment = 1
                        for r in para.runs: r.bold = True
                for row_idx, row in enumerate(rows):
                    for j, val in enumerate(row):
                        table.rows[row_idx+1].cells[j].text = str(val)
                doc.add_paragraph("")
                
        elif b_type == "image":
            image_blocks = [block]
            while i + 1 < len(blocks) and isinstance(blocks[i+1], dict) and blocks[i+1].get("type") == "image":
                image_blocks.append(blocks[i+1])
                i += 1
                
            for j in range(0, len(image_blocks), 2):
                pair = image_blocks[j:j+2]
                if len(pair) == 1:
                    img_url = pair[0].get("url", "")
                    if img_url:
                        local_img = download_file(img_url, temp_folder)
                        if local_img:
                            p = doc.add_paragraph(); p.alignment = 1
                            run = p.add_run(); add_picture_smart(run, local_img, in_table=False)
                            caption = pair[0].get("caption", "")
                            if caption:
                                p_cap = doc.add_paragraph(); p_cap.alignment = 1
                                p_cap.style = get_best_style(doc, "Caption", "Normal").name
                                run_cap = p_cap.add_run(caption); run_cap.font.italic = True
                else:
                    table = doc.add_table(rows=2, cols=2)
                    for col_idx, img_block in enumerate(pair):
                        img_url = img_block.get("url", "")
                        if img_url:
                            local_img = download_file(img_url, temp_folder)
                            if local_img:
                                cell_img = table.cell(0, col_idx)
                                p = cell_img.paragraphs[0]; p.alignment = 1
                                run = p.add_run(); add_picture_smart(run, local_img, in_table=True)
                                caption = img_block.get("caption", "")
                                if caption:
                                    cell_cap = table.cell(1, col_idx)
                                    p_cap = cell_cap.paragraphs[0]; p_cap.alignment = 1
                                    p_cap.style = get_best_style(doc, "Caption", "Normal").name
                                    run_cap = p_cap.add_run(caption); run_cap.font.italic = True
                    doc.add_paragraph("")
        i += 1
    doc.save(output_path)

@app.route('/api/status', methods=['GET'])
def get_status():
    return jsonify({
        "status": "ok",
        "cloud": True,
        "api_key_configured": bool(os.getenv("GEMINI_API_KEY")),
        "version": "2.0-cloud-stateless"
    })

@app.route('/estimate', methods=['POST'])
@require_auth
def estimate_cost():
    try:
        excel_files = request.files.getlist('excel_files')
        template_files = request.files.getlist('template_files')
        source_files = request.files.getlist('source_files')
        knowledge_rules = request.form.get('knowledge_rules', '')
        
        media_count = 0
        text_size_bytes = len(knowledge_rules)
        run_folder = tempfile.mkdtemp(dir=UPLOAD_FOLDER)
        
        try:
            for idx, excel in enumerate(excel_files):
                path = os.path.join(run_folder, f"est_{idx}.xlsx")
                excel.save(path)
                try:
                    df = pd.read_excel(path)
                    for val in df.values.flatten():
                        if isinstance(val, str):
                            urls = extract_urls(val)
                            media_count += len([u for u in urls if 'supabase.co' in u])
                except: pass

            for t in template_files:
                path = os.path.join(run_folder, getattr(t, 'filename', f't_{uuid.uuid4().hex}'))
                t.save(path)
                text_size_bytes += len(extract_text(path))
                    
            for s in source_files:
                path = os.path.join(run_folder, getattr(s, 'filename', f's_{uuid.uuid4().hex}'))
                s.save(path)
                text_size_bytes += len(extract_text(path))

        finally:
            # Força remoção ignorando erros de arquivo em uso no Windows
            shutil.rmtree(run_folder, ignore_errors=True)
            
        text_tokens = int(text_size_bytes / 4)
        total_tokens = text_tokens + (media_count * 258)
        return jsonify({
            "media_count": media_count,
            "text_tokens": text_tokens,
            "total_tokens": total_tokens,
            "estimated_usd": (total_tokens / 1000000) * 1.25
        })
    except Exception as e: return jsonify({"error": str(e)}), 500

@app.route('/prepare_media', methods=['POST'])
@require_auth
def prepare_media():
    try:
        excel_files = request.files.getlist('excel_files')
        folder_id = uuid.uuid4().hex
        run_folder = os.path.join(UPLOAD_FOLDER, folder_id)
        os.makedirs(run_folder, exist_ok=True)
        
        saved_paths = []
        for idx, excel in enumerate(excel_files):
            path = os.path.join(run_folder, f"data_{idx}.xlsx")
            excel.save(path)
            saved_paths.append(path)
            
        def prepare_stream():
            try:
                unique_urls = set()
                yield f"data: {json.dumps({'status': 'Lendo planilhas e extraindo links de mídias...', 'step': 1})}\n\n"
                
                for path in saved_paths:
                    try:
                        df = pd.read_excel(path)
                        for val in df.values.flatten():
                            if isinstance(val, str):
                                for url in extract_urls(val):
                                    if 'supabase.co' in url:
                                        unique_urls.add(url)
                    except: pass
                
                url_to_local = {}
                total_urls = len(unique_urls)
                
                if total_urls > 0:
                    yield f"data: {json.dumps({'status': f'Iniciando download de {total_urls} mídias únicas do servidor...', 'step': 1})}\n\n"
                
                downloaded = 0
                for url in unique_urls:
                    downloaded += 1
                    # Enviar ping a cada download para evitar Timeout no Nginx (Error Network)
                    yield f"data: {json.dumps({'status': f'Baixando mídia {downloaded}/{total_urls}...', 'step': 1})}\n\n"
                    local = download_file(url, run_folder)
                    if local: url_to_local[url] = local
                
                gemini_names = []
                media_mapping_dict = {}
                
                media_files = [local for local in url_to_local.values() if local.lower().endswith(('.jpeg', '.jpg', '.png', '.webp', '.mp3', '.m4a', '.wav', '.ogg', '.mp4', '.avi', '.mov'))]
                total_media = len(media_files)
                processed = 0
                
                if total_media > 0:
                    yield f"data: {json.dumps({'status': f'Iniciando envio de {total_media} mídias para a IA...', 'step': 2})}\n\n"
                    
                for url, local in url_to_local.items():
                    if local.lower().endswith(('.jpeg', '.jpg', '.png', '.webp', '.mp3', '.m4a', '.wav', '.ogg', '.mp4', '.avi', '.mov')):
                        processed += 1
                        try:
                            filename = os.path.basename(local)
                            friendly_name = os.path.basename(url).split('?')[0][-20:]
                            yield f"data: {json.dumps({'status': f'Enviando mídia {processed}/{total_media} ({friendly_name})...', 'step': 2})}\n\n"
                            
                            g_file = client.files.upload(file=local)
                            while g_file.state.name == "PROCESSING":
                                time.sleep(2)
                                g_file = client.files.get(name=g_file.name)
                                yield f"data: {json.dumps({'status': f'Aguardando processamento ({friendly_name})...', 'step': 2})}\n\n"
                            
                            if g_file.state.name == "ACTIVE":
                                gemini_names.append(g_file.name)
                                media_mapping_dict[g_file.name] = url
                                yield f"data: {json.dumps({'status': f'Mídia {processed}/{total_media} pronta!', 'step': 2})}\n\n"
                            else:
                                yield f"data: {json.dumps({'status': f'Aviso: A mídia {friendly_name} falhou (status: {g_file.state.name}).', 'step': 2})}\n\n"
                                
                        except Exception as e:
                            yield f"data: {json.dumps({'status': f'Aviso: Erro ao enviar {friendly_name} ({str(e)}).', 'step': 2})}\n\n"

                yield f"data: {json.dumps({'status': 'Preparo de mídias concluído!', 'step': 3, 'gemini_names': gemini_names, 'media_mapping': json.dumps(media_mapping_dict)})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'error': f'Erro ao preparar mídias: {str(e)}'})}\n\n"
            finally:
                shutil.rmtree(run_folder, ignore_errors=True)
        
        response = Response(stream_with_context(prepare_stream()), mimetype='text/event-stream')
        response.headers['X-Accel-Buffering'] = 'no'
        response.headers['Cache-Control'] = 'no-cache'
        response.headers['Connection'] = 'keep-alive'
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/generate', methods=['POST'])
@require_auth
def generate_report():
    try:
        excel_files = request.files.getlist('excel_files')
        template_files = request.files.getlist('template_files')
        source_files = request.files.getlist('source_files')
        visual_template = request.files.get('visual_template')
        rules = request.form.get('knowledge_rules', '')
        gemini_names_json = request.form.get('gemini_names', '[]')
        media_mapping = request.form.get('media_mapping', '')
        
        try:
            gemini_names = json.loads(gemini_names_json)
        except:
            gemini_names = []

        folder_id = uuid.uuid4().hex
        run_folder = os.path.join(UPLOAD_FOLDER, folder_id)
        os.makedirs(run_folder, exist_ok=True)

        for idx, f in enumerate(excel_files):
            f.save(os.path.join(run_folder, f"data_{idx}.xlsx"))
        if visual_template:
            visual_template.save(os.path.join(run_folder, "master.docx"))
        for idx, t in enumerate(template_files):
            t.save(os.path.join(run_folder, f"ref_{idx}_{t.filename}"))
        for idx, s in enumerate(source_files):
            s.save(os.path.join(run_folder, f"src_{idx}_{s.filename}"))

        def generate_stream():
            try:
                yield f"data: {json.dumps({'status': 'Iniciando processamento dos arquivos...', 'step': 1})}\n\n"
                
                all_excel_str = ""
                url_to_local = {}
                yield f"data: {json.dumps({'status': 'Lendo planilhas de vistoria...', 'step': 2})}\n\n"
                for idx, f in enumerate(excel_files):
                    path = os.path.join(run_folder, f"data_{idx}.xlsx")
                    try:
                        df = pd.read_excel(path)
                        
                        sample_val = ""
                        for val in df.values.flatten():
                            if isinstance(val, str) and len(val) > 20 and not val.startswith('http'):
                                sample_val = val[:50] + "..."
                                break
                        if sample_val:
                            yield f"data: {json.dumps({'status': f'Analisando apontamento: \"{sample_val}\"', 'step': 2})}\n\n"
                            
                        all_excel_str += f"\n[Arquivo {idx+1}]\n" + df.to_json(orient="records", force_ascii=False)
                    except: pass

                yield f"data: {json.dumps({'status': 'Carregando mídias preparadas...', 'step': 3})}\n\n"
                gemini_files = []
                for idx_g, name in enumerate(gemini_names):
                    try:
                        yield f"data: {json.dumps({'status': f'Verificando mídia {idx_g+1}/{len(gemini_names)}...', 'step': 3})}\n\n"
                        g_file = client.files.get(name=name)
                        if g_file.state.name == "ACTIVE":
                            gemini_files.append(g_file)
                    except: pass

                yield f"data: {json.dumps({'status': 'Mapeando referências e fontes normativas...', 'step': 4})}\n\n"
                all_ref_text = ""
                base_template = os.path.join(run_folder, "master.docx") if visual_template else None
                
                for idx, t in enumerate(template_files):
                    yield f"data: {json.dumps({'status': f'Extraindo padrões do modelo: {t.filename}', 'step': 4})}\n\n"
                    path = os.path.join(run_folder, f"ref_{idx}_{t.filename}")
                    if not base_template and path.lower().endswith('.docx'): base_template = path
                    all_ref_text += f"\n[Referência {idx+1}]\n" + extract_text(path)

                sources_text = ""
                for idx, s in enumerate(source_files):
                    yield f"data: {json.dumps({'status': f'Consultando base técnica: {s.filename}', 'step': 4})}\n\n"
                    path = os.path.join(run_folder, f"src_{idx}_{s.filename}")
                    sources_text += f"\n[Fonte {idx+1}]\n" + extract_text(path)

                yield f"data: {json.dumps({'status': 'Analisando dados com Gemini 2.5 Pro (pode demorar alguns minutos)...', 'step': 5})}\n\n"
                sys_inst = (
                    f"Atuação: Você é um especialista em análise de dados e redação técnica de laudos de vistoria.\n"
                    f"Objetivo: Gerar um laudo de vistoria técnico e preciso, utilizando exclusivamente os dados fornecidos na planilha Excel (DADOS), seguindo estritamente a estrutura formal dos laudos de referência anexados.\n\n"
                    f"Diretrizes de Processamento de Fontes (Regras de Ouro):\n"
                    f"1. Laudos de Referência (PDF/Word): Utilize estes arquivos APENAS como guia de estrutura e estilo. Observe onde ficam os títulos, qual a ordem das seções e o tom da linguagem. PROIBIDO: Extrair qualquer dado, fato, localidade, nome ou diagnóstico destes documentos para o novo relatório.\n"
                    f"2. Dados da Vistoria (Excel): Esta é a sua ÚNICA fonte de conteúdo. Todo e qualquer dado (nomes, datas, medições, descrições de patologias, recomendações) deve vir obrigatoriamente das colunas e linhas desta planilha.\n"
                    f"3. Segregação de Conteúdo: Se uma informação não estiver no Excel, ela não existe. Não tente 'completar' o laudo com informações baseadas nos exemplos e evite alucinações. Sinalize a ausência da informação caso ela seja crítica para o laudo.\n\n"
                    f"Instruções de Redação:\n"
                    f"- Mantenha o tom profissional, técnico e objetivo.\n"
                    f"- Organize o relatório conforme as seções identificadas nos laudos de referência (ex: Introdução, Metodologia, Constatações, Conclusão) e considere as Regras do Usuário.\n"
                    f"- Cada campo preenchido no relatório final deve ter correspondência direta com uma célula do Excel.\n\n"
                    f"Formatos permitidos OBRIGATÓRIOS (Retorne APENAS um ARRAY JSON de objetos):\n"
                    f"1. Texto: {{'type': 'heading1'|'paragraph', 'text': 'conteudo'}}\n"
                    f"2. Tabela: {{'type': 'table', 'headers': [], 'rows': [[]]}}\n"
                    f"3. Imagem: {{'type': 'image', 'url': 'URL_FORNECIDA', 'caption': 'Legenda técnica'}}\n\n"
                    f"REGRAS DE IMAGEM: Não inclua 100% das fotos se redundantes, mas INCLUA TODAS as relevantes para ilustrar anomalias/apontamentos. Toda imagem deve ter legenda descritiva. Use a URL exata do MAPEAMENTO DE IMAGENS.\n"
                    f"IMPORTANTE: NUNCA use tags HTML. Gere apenas texto puro.\n"
                    f"Verificação de Erros Final: Antes de entregar o JSON, faça uma dupla checagem se algum dado veio dos exemplos em vez do Excel. Se sim, apague e substitua pelos dados corretos do Excel ou deixe o campo em branco.\n"
                    f"Regras do Usuário adicionais: {rules}"
                )
                user_prompt = f"REFERÊNCIAS:\n{all_ref_text[:15000]}\n\nFONTES:\n{sources_text[:15000]}\n\nDADOS:\n{all_excel_str}\n"
                
                try:
                    media_mapping_dict = json.loads(media_mapping)
                except:
                    media_mapping_dict = {}
                    
                contents = []
                for g_file in gemini_files:
                    url = media_mapping_dict.get(g_file.name, "")
                    if url:
                        contents.append(f"[ESTA É A IMAGEM REFERENTE À URL: {url}]")
                    contents.append(g_file)
                contents.append(user_prompt)

                # Keep-alive para proxy
                q = queue.Queue()
                def fetch_gemini():
                    try:
                        resp = client.models.generate_content_stream(
                            model='gemini-2.5-pro',
                            contents=contents,
                            config=genai.types.GenerateContentConfig(system_instruction=sys_inst, response_mime_type="application/json")
                        )
                        for chunk in resp:
                            q.put(("chunk", chunk))
                        q.put(("done", None))
                    except Exception as e:
                        q.put(("error", e))

                t = threading.Thread(target=fetch_gemini)
                t.start()

                raw_text = ""
                used_tokens = 0
                while True:
                    try:
                        msg_type, data = q.get(timeout=10) # Acorda a cada 10s
                        if msg_type == "done":
                            break
                        elif msg_type == "error":
                            raise data
                        elif msg_type == "chunk":
                            raw_text += data.text
                            if data.usage_metadata:
                                used_tokens = data.usage_metadata.total_token_count
                            yield f"data: {json.dumps({'status': f'Analisando dados e gerando texto (recebidos {len(raw_text)} bytes)...', 'step': 5})}\n\n"
                    except queue.Empty:
                        # Timeout do queue atingido, envia ping de dados para manter conexao viva e forçar flush
                        yield f"data: {json.dumps({'status': 'Análise da IA em andamento, aguarde...', 'step': 5})}\n\n"
                
                yield f"data: {json.dumps({'status': 'Montando documento Word...', 'step': 6})}\n\n"
                
                raw_text = raw_text.strip()
                if raw_text.startswith("```json"): raw_text = raw_text[7:]
                elif raw_text.startswith("```"): raw_text = raw_text[3:]
                if raw_text.endswith("```"): raw_text = raw_text[:-3]
                
                parsed_json = json.loads(raw_text.strip())
                
                final_doc = os.path.join(run_folder, "Resultado.docx")
                
                q_doc = queue.Queue()
                def build_doc():
                    try:
                        reconstruct_doc(parsed_json, base_template, final_doc, run_folder)
                        
                        final_doc_obj = docx.Document(final_doc)
                        if len(final_doc_obj.paragraphs) <= 1 and not final_doc_obj.tables:
                            raise Exception(f"A IA gerou um formato inesperado que resultou num laudo vazio. Retorno cru da IA (copie isto e envie ao suporte): {raw_text[:2000]}")
                        q_doc.put(("done", None))
                    except Exception as e:
                        q_doc.put(("error", e))

                t_doc = threading.Thread(target=build_doc)
                t_doc.start()

                while True:
                    try:
                        msg_type, doc_data = q_doc.get(timeout=10)
                        if msg_type == "done":
                            break
                        elif msg_type == "error":
                            raise doc_data
                    except queue.Empty:
                        yield ": keep-alive\n\n"
                
                yield f"data: {json.dumps({'status': 'Concluído!', 'step': 7, 'file_id': folder_id, 'tokens_used': used_tokens})}\n\n"
            except Exception as e:
                error_msg = str(e)
                if "403" in error_msg or "PERMISSION_DENIED" in error_msg or "API key" in error_msg or "API_KEY_INVALID" in error_msg:
                    error_msg = "Falha de Autenticação na IA: A chave da API do Gemini atual é inválida, vazou ou foi bloqueada pelo Google. Por favor, gere uma nova chave e atualize no painel do servidor."
                elif "429" in error_msg or "quota" in error_msg.lower() or "exhausted" in error_msg.lower():
                    error_msg = "Limite da IA Excedido: A cota de uso da API do Gemini foi atingida. Tente novamente mais tarde."
                elif "500" in error_msg or "503" in error_msg:
                    error_msg = "IA Indisponível: Os servidores do Google estão instáveis no momento. Tente novamente em alguns minutos."
                else:
                    error_msg = f"Erro ao gerar laudo: {error_msg}"
                yield f"data: {json.dumps({'error': error_msg})}\n\n"

        response = Response(stream_with_context(generate_stream()), mimetype='text/event-stream')
        response.headers['X-Accel-Buffering'] = 'no'
        response.headers['Cache-Control'] = 'no-cache'
        response.headers['Connection'] = 'keep-alive'
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/download/<folder_id>', methods=['GET'])
@require_auth
def download_report(folder_id):
    import re
    if not re.match(r'^[a-zA-Z0-9_-]+$', folder_id):
        return "Invalid folder", 400
    
    final_doc = os.path.join(UPLOAD_FOLDER, folder_id, "Resultado.docx")
    project_name = request.args.get('project_name', 'Laudo_Gerado')
    # Sanitizar nome do arquivo
    safe_name = re.sub(r'[^a-zA-Z0-9_\- ]', '_', project_name).strip().replace(' ', '_')
    download_name = f"{safe_name}_IA.docx"
    
    if os.path.exists(final_doc):
        return send_file(final_doc, as_attachment=True, download_name=download_name)
    return "File not found", 404



@app.route('/api/admin/users', methods=['GET'])
@require_admin
def admin_get_users():
    if not SUPABASE_SERVICE_ROLE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY não configurado no backend."}), 500
    
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}"
    }
    
    response = requests.get(f"{SUPABASE_URL}/auth/v1/admin/users", headers=headers)
    if response.status_code == 200:
        return jsonify(response.json())
    return jsonify({"error": response.text}), response.status_code

@app.route('/api/admin/users', methods=['POST'])
@require_admin
def admin_create_user():
    if not SUPABASE_SERVICE_ROLE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY não configurado no backend."}), 500
        
    data = request.json
    email = data.get('email')
    password = data.get('password')
    name = data.get('name', '')
    
    if not email or not password:
        return jsonify({"error": "E-mail e senha são obrigatórios"}), 400
        
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "email": email,
        "password": password,
        "email_confirm": True, # Auto-confirmar
        "user_metadata": {
            "name": name,
            "raw_password": password
        }
    }
    
    response = requests.post(f"{SUPABASE_URL}/auth/v1/admin/users", headers=headers, json=payload)
    if response.status_code in [200, 201]:
        return jsonify(response.json())
    return jsonify({"error": response.text}), response.status_code

@app.route('/api/admin/users/<user_id>', methods=['DELETE'])
@require_admin
def admin_delete_user(user_id):
    if not SUPABASE_SERVICE_ROLE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY não configurado no backend."}), 500
        
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}"
    }
    
    response = requests.delete(f"{SUPABASE_URL}/auth/v1/admin/users/{user_id}", headers=headers)
    if response.status_code in [200, 204]:
        return jsonify({"success": True})
    return jsonify({"error": response.text}), response.status_code

@app.route('/api/admin/users/<user_id>/reset', methods=['PUT'])
@require_admin
def admin_reset_user_password(user_id):
    if not SUPABASE_SERVICE_ROLE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY não configurado no backend."}), 500
        
    data = request.json
    new_password = data.get('password')
    if not new_password:
        return jsonify({"error": "A nova senha é obrigatória"}), 400
        
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json"
    }
    
    # 1. Obter os metadados atuais do usuário
    get_res = requests.get(f"{SUPABASE_URL}/auth/v1/admin/users/{user_id}", headers=headers)
    if get_res.status_code != 200:
        return jsonify({"error": "Usuário não encontrado"}), 404
        
    user_data = get_res.json()
    metadata = user_data.get("user_metadata", {})
    if "must_change_password" in metadata:
        del metadata["must_change_password"]
    metadata["raw_password"] = new_password
    
    payload = {
        "password": new_password,
        "user_metadata": metadata
    }
    
    # 2. Atualizar a senha e os metadados
    response = requests.put(f"{SUPABASE_URL}/auth/v1/admin/users/{user_id}", headers=headers, json=payload)
    if response.status_code in [200, 204]:
        return jsonify({"success": True})
    return jsonify({"error": response.text}), response.status_code

@app.route('/api/admin/users/reset-tokens', methods=['PUT'])
@require_admin
def admin_reset_all_tokens():
    if not SUPABASE_SERVICE_ROLE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY não configurado no backend."}), 500
        
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json"
    }
    
    response = requests.get(f"{SUPABASE_URL}/auth/v1/admin/users", headers=headers)
    if response.status_code != 200:
        return jsonify({"error": response.text}), response.status_code
        
    users_data = response.json()
    users = users_data.get('users', users_data) if isinstance(users_data, dict) else users_data
    if not isinstance(users, list):
        return jsonify({"error": "Formato de usuários inválido"}), 500

    updated = 0
    for user in users:
        metadata = user.get("user_metadata", {})
        if metadata.get("total_tokens_used", 0) > 0:
            metadata["total_tokens_used"] = 0
            payload = {"user_metadata": metadata}
            update_res = requests.put(f"{SUPABASE_URL}/auth/v1/admin/users/{user['id']}", headers=headers, json=payload)
            if update_res.status_code in [200, 204]:
                updated += 1

    return jsonify({"success": True, "updated_count": updated})

@app.route('/generate_photo_report_stream', methods=['POST'])
def generate_photo_report_stream():
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    q = queue.Queue()
    
    def process_report():
        run_id = str(uuid.uuid4())
        run_folder = os.path.join(UPLOAD_FOLDER, f"photo_report_{run_id}")
        os.makedirs(run_folder, exist_ok=True)
        
        try:
            q.put({"status": "Iniciando geração do relatório fotográfico...", "step": 1})
            
            doc = Document()
            
            # Cabeçalho
            title = data.get("title", "Vistoria")
            address = data.get("address", "")
            engineer = data.get("engineer", "")
            
            heading = doc.add_heading(f"Relatório de Vistoria: {title}", level=1)
            heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            if address:
                p_add = doc.add_paragraph(f"Endereço: {address}")
                p_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                
            p_eng = doc.add_paragraph(f"Engenheiro/Responsável: {engineer}")
            p_eng.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            p_date = doc.add_paragraph(f"Relatório gerado em: {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M:%S')}")
            p_date.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            records = data.get("records", [])
            total_records = len(records)
            
            # Fonte Arial 12
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = docx.shared.Pt(12)
            
            for idx, record in enumerate(records):
                q.put({"status": f"Processando registro {idx+1}/{total_records}...", "step": 2})
                
                doc.add_heading(record.get("form_title", f"Registro #{idx+1}"), level=2)
                p_meta = doc.add_paragraph(f"Lançado por {record.get('inspector', '')} em {record.get('date', '')}")
                p_meta.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)
                
                text_data = record.get("text_data", {})
                for k, v in text_data.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k}: ").bold = True
                    p.add_run(str(v))
                    
                images = record.get("images", [])
                audios = record.get("audios", [])
                
                # Baixar imagens e corrigir orientação
                downloaded_images = []
                for i_idx, img_url in enumerate(images):
                    try:
                        q.put({"status": f"Baixando imagem {i_idx+1} do registro {idx+1}...", "step": 3})
                        res = requests.get(img_url, timeout=15)
                        if res.status_code == 200:
                            img_path = os.path.join(run_folder, f"img_{idx}_{i_idx}.jpg")
                            with open(img_path, 'wb') as f:
                                f.write(res.content)
                            fix_image_orientation(img_path)
                            
                            # Checar orientacao para tabela
                            with Image.open(img_path) as im:
                                w, h = im.size
                                is_landscape = w >= h
                            downloaded_images.append({"path": img_path, "landscape": is_landscape})
                    except Exception as e:
                        print(f"Erro ao baixar imagem: {e}")
                        
                # Montar tabela de imagens no Word
                if downloaded_images:
                    # Agrupar imagens em linhas: paisagem=1 por linha, retrato=2 por linha
                    rows = []
                    current_row = []
                    for img in downloaded_images:
                        if img["landscape"]:
                            if current_row:
                                rows.append(current_row)
                                current_row = []
                            rows.append([img])
                        else:
                            current_row.append(img)
                            if len(current_row) == 2:
                                rows.append(current_row)
                                current_row = []
                    if current_row:
                        rows.append(current_row)
                        
                    table = doc.add_table(rows=len(rows), cols=2)
                    for r_idx, row_imgs in enumerate(rows):
                        if len(row_imgs) == 1 and row_imgs[0]["landscape"]:
                            # Mesclar colunas para paisagem
                            cell = table.cell(r_idx, 0)
                            cell.merge(table.cell(r_idx, 1))
                            p = cell.paragraphs[0]
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(row_imgs[0]["path"], width=Cm(15))
                        else:
                            for c_idx, img in enumerate(row_imgs):
                                cell = table.cell(r_idx, c_idx)
                                p = cell.paragraphs[0]
                                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run()
                                run.add_picture(img["path"], width=Cm(7.5))
                                
                # IA Legenda
                if (images or audios) and api_key:
                    q.put({"status": f"Gerando legenda com IA (Gemini 2.5 Flash) para o registro {idx+1}...", "step": 4})
                    try:
                        contents = [f"DADOS DO REGISTRO:\n{json.dumps(text_data, ensure_ascii=False)}"]
                        
                        uploaded_gemini = []
                        for img in downloaded_images:
                            g_file = client.files.upload(file=img["path"])
                            uploaded_gemini.append(g_file)
                            contents.append(g_file)
                            
                        for a_idx, aud_url in enumerate(audios):
                            res = requests.get(aud_url, timeout=15)
                            if res.status_code == 200:
                                aud_path = os.path.join(run_folder, f"aud_{idx}_{a_idx}.m4a")
                                with open(aud_path, 'wb') as f:
                                    f.write(res.content)
                                g_file = client.files.upload(file=aud_path)
                                uploaded_gemini.append(g_file)
                                contents.append(g_file)
                                
                        contents.append("Crie UMA ÚNICA legenda consolidada, técnica e descritiva para este conjunto de imagens, incorporando as descrições dos áudios e os dados de texto. Seja direto, evite floreios e NÃO pule linhas. Máximo 3 frases.")
                        
                        resp = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=contents
                        )
                        caption = resp.text.strip()
                        
                        # Limpar arquivos do Gemini
                        for g_f in uploaded_gemini:
                            try:
                                client.files.delete(name=g_f.name)
                            except:
                                pass
                                
                    except Exception as e:
                        caption = f"[Erro ao gerar legenda IA: {str(e)}]"
                else:
                    caption = "Registros fotográficos do apontamento."
                    
                if downloaded_images:
                    p_cap = doc.add_paragraph(caption)
                    p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].italic = True
                    
            output_filename = f"Relatorio_Fotografico_{run_id}.docx"
            output_path = os.path.join(UPLOAD_FOLDER, output_filename)
            doc.save(output_path)
            
            q.put({"status": "Concluído", "download_url": f"/download_photo_report/{output_filename}", "step": 5})
            
        except Exception as e:
            q.put({"error": f"Erro interno: {str(e)}"})
        finally:
            q.put(None)
            
    threading.Thread(target=process_report).start()
    
    def generate():
        while True:
            msg = q.get()
            if msg is None:
                break
            yield f"data: {json.dumps(msg)}\n\n"
            
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/download_photo_report/<filename>', methods=['GET'])
def download_photo_report(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({"error": "Arquivo não encontrado"}), 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
